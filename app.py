import streamlit as st
from datetime import datetime, date
from PIL import Image as PILImage
import numpy as np
import io
from docx import Document
from docx.shared import Inches
from streamlit_drawable_canvas import st_canvas
import smtplib
from email.message import EmailMessage
import shutil
import re


# Set page configuration with a favicon
st.set_page_config(
    page_title="Prevista Skills Bootcamp",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png",  # Path to your logo
    layout="centered"  # "centered" or "wide"
)

def update_doc():
    pass

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = ", ".join(receiver_email)
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))


def replace_placeholders(template_file, modified_file, placeholder_values, signature_path):
    try:
        print(f"Copying template file '{template_file}' to '{modified_file}'...")
        shutil.copy(template_file, modified_file)

        print(f"Opening document '{modified_file}'...")
        doc = Document(modified_file)

        # Function to convert value to string, handling datetime.date objects
        def convert_to_str(value):
            if isinstance(value, date):
                return value.strftime('%Y-%m-%d')  # Convert date to string
            return str(value)  # Convert other types to string

        # Compile regular expressions for all placeholders
        placeholders = {re.escape(key): convert_to_str(value) for key, value in placeholder_values.items()}
        placeholders_pattern = re.compile(r'\b(' + '|'.join(placeholders.keys()) + r')\b')

        # Replace placeholders in paragraphs
        print("Replacing placeholders in paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
            if original_text != updated_text:
                print(f"Updated paragraph text: '{original_text}' -> '{updated_text}'")
                para.text = updated_text

        # Replace placeholders in tables
        print("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
                        if original_text != updated_text:
                            print(f"Updated table cell text: '{original_text}' -> '{updated_text}'")
                            para.text = updated_text

                    # Inspect cell runs
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            run_updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], run_text)
                            if run_text != run_updated_text:
                                print(f"Updated run text in table cell: '{run_text}' -> '{run_updated_text}'")
                                run.text = run_updated_text

        # Check and handle signature placeholder
        print("Inspecting document for 'ph_signature' placeholder...")
        signature_placeholder_found = False

        # Check paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()  # Remove any extra spaces around text
            while 'ph_signature' in para_text:
                print(f"Found 'ph_signature' in paragraph: '{para_text}'")
                para_text = para_text.replace('ph_signature', '').strip()  # Remove 'ph_signature' and any leading/trailing spaces
                para.text = para_text
                resized_image_path = 'resized_signature_image.png'
                
                try:
                    # Open and resize the image
                    print(f"Opening image file: {signature_path}")
                    resized_image = PILImage.open(signature_path)
                    print(f"Original image size: {resized_image.size}")
                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                    resized_image.save(resized_image_path)  # Save resized image to a file
                    print(f"Resized image saved to: {resized_image_path}")
                    
                    # Add picture to the paragraph
                    print(f"Adding picture to paragraph from path: {resized_image_path}")
                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                    print("Inserted signature image into paragraph.")
                    signature_placeholder_found = True
                except Exception as img_e:
                    print(f"An error occurred with image processing: {img_e}")

        # Check table cells again in case the placeholder was missed
        if not signature_placeholder_found:
            print("Checking table cells for 'ph_signature'...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            while 'ph_signature' in para_text:
                                print(f"Found 'ph_signature' in table cell paragraph: '{para_text}'")
                                para_text = para_text.replace('ph_signature', '').strip()
                                para.text = para_text
                                resized_image_path = 'resized_signature_image.png'
                                
                                try:
                                    # Open and resize the image
                                    print(f"Opening image file: {signature_path}")
                                    resized_image = PILImage.open(signature_path)
                                    print(f"Original image size: {resized_image.size}")
                                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                                    resized_image.save(resized_image_path)  # Save resized image to a file
                                    print(f"Resized image saved to: {resized_image_path}")
                                    
                                    # Add picture to the table cell
                                    print(f"Adding picture to table cell from path: {resized_image_path}")
                                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                                    print("Inserted signature image into table cell.")
                                    signature_placeholder_found = True
                                except Exception as img_e:
                                    print(f"An error occurred with image processing: {img_e}")

        if not signature_placeholder_found:
            print("No signature placeholder found.")

        # Save the modified document
        print(f"Saving modified document '{modified_file}'...")
        doc.save(modified_file)
        print(f"Document modification complete: '{modified_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")

    # file download button
    with open(modified_file, 'rb') as f:
        file_contents = f.read()
        st.download_button(
            label="Download Your Response",
            data=file_contents,
            file_name=modified_file,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )


if 'files' not in st.session_state:
    st.session_state.files = []

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
    st.session_state.first_name = ""
    st.session_state.sir_name = ""
    st.session_state.dob = None
    st.session_state.address = ""
    st.session_state.city = ""
    st.session_state.postcode = ""
    st.session_state.phone_number = ""
    st.session_state.email = ""
    st.session_state.ni_number = ""
    st.session_state.emergency_contact_name = ""
    st.session_state.emergency_contact_relationship = ""
    st.session_state.emergency_contact_phone = ""
    st.session_state.emergency_contact_email = ""
    st.session_state.highest_education = ""
    st.session_state.institution_name = ""
    st.session_state.year_of_completion = ""
    st.session_state.subject=""
    st.session_state.date_applied=""
    st.session_state.current_employment_status = "Unemployed"

    st.session_state.current_employer_name = ""
    st.session_state.current_employer_postcode = ""
    st.session_state.applied_through_employer = ""
    st.session_state.hours_worked_per_week = ""
    st.session_state.estimated_salary = ""
    st.session_state.salary_frequency = ""
    st.session_state.contract_type = ""
    st.session_state.continue_working = ""
    st.session_state.main_work = ""
    st.session_state.previous_employment_job_title = ""
    st.session_state.previous_employment_industry = ""

    # st.session_state.previous_employment_employer_name = ''
    # st.session_state.previous_employment_job_title = ''
    # st.session_state.previous_employment_duration = ''
    st.session_state.disabilities = ""
    st.session_state.benefits = ""
    st.session_state.benefits_details = ""
    st.session_state.statement_of_interest = ""
    st.session_state.career_goals = ""
    st.session_state.preferred_industry = ""
    st.session_state.support_needed = ""
    st.session_state.age_verified = False
    st.session_state.residency_verified = False
    # st.session_state.unemployment_verified = False
    st.session_state.education_verified = False
    st.session_state.commitment_verified = False
    st.session_state.documents_verified = False
    st.session_state.submission_done = False

    st.session_state.completed_level_6_or_above=''
    st.session_state.claiming_universal_credit = ''
    st.session_state.caring_responsibilities = ''
    st.session_state.gender = ''
    st.session_state.disabilities_health_condition = ''
    st.session_state.ethnicity = ''

# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)




# Define the total number of steps
total_steps = 11
# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)
# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)


# Define different steps
if st.session_state.step == 1:
    st.image('resources/header-wihout-bg.png', use_column_width=True)

    st.title("Skills Bootcamp Enrollment and Registration Document")
    st.write("Provider: Prevista Ltd. | Sponsor: Surrey County Council | Website: www.prevista.co.uk")
    st.write("________________________________________")
    st.write("**Application and Enrollment Form**")
    st.write("""
    Please complete the upcoming sections to finalize your enrollment.
    """)
    if st.button("Next"):
        # st.session_state.step = 2
        st.session_state.step = 10
        st.experimental_rerun()

elif st.session_state.step == 2:
    st.title("> 1: Personal Information")
    st.session_state.first_name = st.text_input("First Name")
    st.session_state.sir_name = st.text_input("Sir Name")
    st.session_state.dob = st.date_input("Date of Birth", 
                                        min_value=date(1900, 1, 1),  # Minimum selectable date
                                        max_value=date.today(),  # Maximum selectable date
                                        key="date_of_borth",  # Unique key for the widget
                                        help="Choose a date",  # Tooltip text
                                        value=st.session_state.dob or datetime(2000, 1, 1), 
                                        format='DD/MM/YYYY')
    
    st.session_state.gender = st.selectbox("APPLICANT’S GENDER", ["Male", "Female", "Other", "Prefer not to say"])
    st.session_state.disabilities_health_condition = st.radio("DOES APPLICANT HAVE DISABILITIES OR LONG TERM HEALTH CONDITION?", ["No", "Yes"])
    st.session_state.ethnicity = st.selectbox("APPLICANT’S ETHNICITY", [
        "White",
        "Mixed/Multiple ethnic groups",
        "Asian/Asian British",
        "Black/African/Caribbean/Black British",
        "Other ethnic group",
        "Prefer not to say"
    ])
    
    st.session_state.address = st.text_input("Address")
    st.session_state.city = st.text_input("City")
    st.session_state.postcode = st.text_input("Home Postcode")
    st.session_state.phone_number = st.text_input("Phone Number")
    st.session_state.email = st.text_input("Email Address")
    st.session_state.ni_number = st.text_input("National Insurance Number")
    st.session_state.emergency_contact_name = st.text_input("Emergency Contact Full Name")
    st.session_state.emergency_contact_relationship = st.text_input("Emergency Contact Relationship")
    st.session_state.emergency_contact_phone = st.text_input("Emergency Contact Phone Number")
    st.session_state.emergency_contact_email = st.text_input("Emergency Contact Email Address")
    
    if st.button("Next"):
        if (st.session_state.first_name and st.session_state.sir_name and st.session_state.dob and 
            st.session_state.address and st.session_state.city and
            st.session_state.postcode and st.session_state.phone_number and
            st.session_state.email and st.session_state.ni_number and
            st.session_state.emergency_contact_name and st.session_state.emergency_contact_relationship and
            st.session_state.gender and st.session_state.disabilities_health_condition and st.session_state.ethnicity and
            st.session_state.emergency_contact_phone and st.session_state.emergency_contact_email):
            st.session_state.step = 3
            st.experimental_rerun()
        else:
            st.warning("Please fill in all fields before proceeding.")

elif st.session_state.step == 3:
    st.title("> 2: Educational Background")
    st.session_state.highest_education = st.text_input("Highest Level of Education")
    st.session_state.institution_name = st.text_input("Institution Name")
    st.session_state.year_of_completion = st.text_input("Year of Completion")

    st.session_state.completed_level_6_or_above = st.radio(
        "COMPLETED LEVEL 6 OR ABOVE?", 
        options=["Yes", "No"]
    )
    if st.session_state.completed_level_6_or_above == "Yes":
        st.session_state.subject = st.text_input("IF YES, WHAT SUBJECT")
    else:
        st.session_state.subject = ""

    st.session_state.date_applied = st.text_input("DATE APPLIED TO JOIN SKILLS BOOTCAMP? (MM/YYYY)")

    if st.button("Next"):
        if (st.session_state.highest_education and 
            st.session_state.institution_name and 
            st.session_state.year_of_completion and 
            st.session_state.date_applied):
            
            # Check if level 6 or above is completed and if subject field is filled
            if st.session_state.completed_level_6_or_above == "Yes" and not st.session_state.subject:
                st.warning("Please provide the subject if you completed level 6 or above.")
            else:
                # Proceed to the next step
                st.session_state.step = 4
                st.experimental_rerun()
        else:
            st.warning("Please provide all required educational background details before proceeding.")


elif st.session_state.step == 4:
    st.title("> 3: Employment Status")
    
    # Current employment status
    st.session_state.current_employment_status = st.radio(
        "APPLICANT’S EMPLOYMENT STATUS BEFORE JOINING SKILLS BOOTCAMP", 
        ["Unemployed", "Employed"]
    )

    if st.session_state.current_employment_status == "Employed":
        st.session_state.current_employer_name = st.text_input("NAME OF CURRENT EMPLOYER")
        st.session_state.current_employer_postcode = st.text_input("POSTCODE OF CURRENT EMPLOYER")
        st.session_state.applied_through_employer = st.radio("DID YOU APPLIED THROUGH CURRENT EMPLOYER?", ["Yes", "No"])
        st.session_state.hours_worked_per_week = st.number_input("PRIOR TO APPLYING TO JOIN SKILLS BOOTCAMP, HOW MANY HOURS WORKED PER WEEK IN THEIR JOB(S)?", min_value=0, step=1)
        st.session_state.estimated_salary = st.number_input("ESTIMATED SALARY (GBP)", min_value=0.0, step=0.01)
        st.session_state.salary_frequency = st.selectbox("SALARY PAID HOURLY/WEEKLY/MONTHLY/YEARLY…?", ["Hourly", "Weekly", "Monthly", "Yearly"])
        st.session_state.contract_type = st.selectbox("TYPE OF CONTRACT?", ["Full Time", "Permanent", "Temporary", "Part-Time", "Zero Hour"])
        st.session_state.continue_working = st.radio("ARE YOU GOING TO CONTINUE WORKING WHILE ON THE SKILLS BOOTCAMP?", ["Yes", "No"])
        st.session_state.main_work = st.text_input("YOUR MAIN WORK PRIOR TO JOINING THE SKILLS BOOTCAMP?")
    else:
        # Fields specific to unemployed applicants
        st.session_state.previous_employment_job_title = st.text_input("IF CURRENTLY UNEMPLOYED, WHAT IS THE MOST RECENT WORK? JOB TITLE")
        st.session_state.previous_employment_industry = st.text_input("WHAT IS THE INDUSTRY OF MAIN WORK OR MOST RECENT WORK?")

    # Additional fields for all applicants
    st.session_state.claiming_universal_credit = st.radio("ARE YOU CURRENTLY CLAIMING UNIVERSAL CREDIT?", ["Yes", "No"])
    st.session_state.caring_responsibilities = st.radio("DO YOU HAVE CHILDREN OR ADULTS CARING RESPONSIBILITIES?", ["Yes", "No"])


    # Form submission
    if st.button("Next"):
        # Basic validation for required fields
        required_fields = [
            # st.session_state.date_of_birth,
            # st.session_state.gender,
            # st.session_state.ethnicity
        ]
        
        if st.session_state.current_employment_status == "Employed":
            required_fields.extend([
                st.session_state.current_employer_name,
                st.session_state.current_employer_postcode,
                st.session_state.hours_worked_per_week,
                st.session_state.estimated_salary,
                st.session_state.salary_frequency,
                st.session_state.contract_type
            ])
        else:
            required_fields.append(st.session_state.previous_employment_job_title)
        
        # Check if any required field is empty or not filled in
        if all(required_fields):
            # Proceed to the next step
            st.session_state.step = 6
            st.experimental_rerun()
        else:
            st.warning("Please provide all required details before proceeding.")


elif st.session_state.step == 5:
    st.title("> 4: Additional Information")
    st.session_state.disabilities = st.selectbox("Do you have any disabilities or learning difficulties?", ["Yes", "No"])
    if st.session_state.disabilities == "Yes":
        st.session_state.disabilities_details = st.text_area("If yes, please provide details:")
    st.session_state.benefits = st.selectbox("Are you receiving any benefits?", ["Yes", "No"])
    if st.session_state.benefits == "Yes":
        st.session_state.benefits_details = st.text_area("If yes, please specify:")
    
    if st.button("Next"):
        st.session_state.step = 6
        st.experimental_rerun()

elif st.session_state.step == 6:
    st.title("> 5: Statement of Interest")
    st.session_state.statement_of_interest = st.text_area("Please explain why you are interested in the Skills Bootcamp and how you think it will benefit you (200 words):")
    
    # Check word count and provide feedback
    if len(st.session_state.statement_of_interest.split()) > 200:
        st.warning("Your statement exceeds the 200-word limit. Please shorten it.")
    
    if st.button("Next"):
        if st.session_state.statement_of_interest:
            st.session_state.step = 7
            st.experimental_rerun()
        else:
            st.warning("Please provide a statement of interest before proceeding.")

elif st.session_state.step == 7:
    st.title("> 6: CEIAG (Careers Education, Information, Advice, and Guidance)")
    st.session_state.career_goals = st.text_area("Career Goals")
    st.session_state.preferred_industry = st.text_input("Preferred Industry")
    st.session_state.support_needed = st.text_area("Support Needed (e.g., resume writing, interview skills)")
    
    if st.button("Next"):
        st.session_state.step = 8
        st.experimental_rerun()

elif st.session_state.step == 8:
    st.title("> 7: Eligibility Criteria")
    st.write("Eligibility Requirements:")
    st.write("1. Age: Applicants must be 19 years or older.")
    # st.write("2. Employment Status: Applicants must be unemployed at the time of enrollment.")
    st.write("2. Residency: Applicants must be residents of Surrey.")
    st.write("3. Education: Applicants must have at least a basic level of literacy and numeracy.")
    st.write("4. Commitment: Applicants must be able to commit to the full duration of the bootcamp (12 weeks, 3 days per week).")
    st.session_state.age_verified = st.checkbox("I confirm I am 19 years or older")
    st.session_state.residency_verified = st.checkbox("I confirm I am a resident of Surrey")
    # st.session_state.unemployment_verified = st.checkbox("I confirm I am currently unemployed")
    st.session_state.education_verified = st.checkbox("I confirm I have at least a basic level of literacy and numeracy")
    st.session_state.commitment_verified = st.checkbox("I confirm I can commit to the full duration of the bootcamp (12 weeks, 3 days per week).")
    
    if st.button("Next"):
        if (st.session_state.age_verified and st.session_state.residency_verified and
            # st.session_state.unemployment_verified and 
            st.session_state.education_verified and st.session_state.commitment_verified):
            st.session_state.step = 9
            st.experimental_rerun()
        else:
            st.warning("Please verify all eligibility criteria before proceeding.")

elif st.session_state.step == 9:
    st.title("> 8: Supporting Documents")
    
    st.write("Please upload any relevant documents if you have them:")
    
    # File uploaders for each document

    # # st.session_state.files["proof_of_age_front"] = st.file_uploader("Proof of Age (Front)", type=["jpg", "png", "pdf"], key="proof_of_age_front")
    st.session_state.proof_of_age_front = st.file_uploader("Proof of Age (E.G., PASSPORT, DRIVING LICENSE) [Front]", type=["jpg", "png", "pdf", "docx"], key="proof_of_age_front_key")
    if st.session_state.proof_of_age_front is not None:
        if st.session_state.proof_of_age_front not in st.session_state.files:
            st.session_state.files.append(st.session_state.proof_of_age_front)

    # # st.session_state.files["proof_of_age_back"] = st.file_uploader("Proof of Age (Back)", type=["jpg", "png", "pdf"], key="proof_of_age_back")
    st.session_state.proof_of_age_back = st.file_uploader("Proof of Age (E.G., PASSPORT, DRIVING LICENSE) [Back]", type=["jpg", "png", "pdf", "docx"], key="proof_of_age_back_key")
    if st.session_state.proof_of_age_back is not None:
        if st.session_state.proof_of_age_back not in st.session_state.files:
            st.session_state.files.append(st.session_state.proof_of_age_back)


    # # st.session_state.files["proof_of_residency_front"] = st.file_uploader("Proof of Residency (Front)", type=["jpg", "png", "pdf"], key="proof_of_residency_front")
    st.session_state.proof_of_residency_front = st.file_uploader("Proof of Residency (E.G., UTILITY BILL, COUNCIL TAX STATEMENT) [Front]", type=["jpg", "png", "pdf", "docx"], key="proof_of_residency_front_key")
    if st.session_state.proof_of_residency_front is not None:
        if st.session_state.proof_of_residency_front not in st.session_state.files:
            st.session_state.files.append(st.session_state.proof_of_residency_front)

    # # st.session_state.files["proof_of_residency_back"] = st.file_uploader("Proof of Residency (Back)", type=["jpg", "png", "pdf"], key="proof_of_residency_back")
    st.session_state.proof_of_residency_back = st.file_uploader("Proof of Residency (E.G., UTILITY BILL, COUNCIL TAX STATEMENT) [Back]", type=["jpg", "png", "pdf", "docx"], key="proof_of_residency_back_key")
    if st.session_state.proof_of_residency_back is not None:
        if st.session_state.proof_of_residency_back not in st.session_state.files:
            st.session_state.files.append(st.session_state.proof_of_residency_back)
    

    # # st.session_state.files["proof_of_unemployment_front"] = st.file_uploader("Proof of Unemployment (Front)", type=["jpg", "png", "pdf"], key="proof_of_unemployment_front")
    st.session_state.proof_of_unemployment_front = st.file_uploader("Proof of Unemployment (E.G., BENEFIT STATEMENT, LETTER FROM JOBCENTRE)", type=["jpg", "png", "pdf", "docx"], key="proof_of_unemployment_front_key")
    if st.session_state.proof_of_unemployment_front is not None:
        if st.session_state.proof_of_unemployment_front not in st.session_state.files:
            st.session_state.files.append(st.session_state.proof_of_unemployment_front)
            
    # # st.session_state.files["proof_of_unemployment_back"] = st.file_uploader("Proof of Unemployment (Back)", type=["jpg", "png", "pdf"], key="proof_of_unemployment_back")
    # st.session_state.proof_of_unemployment_back = st.file_uploader("Proof of Unemployment (Back)", type=["jpg", "png", "pdf", "docx"], key="proof_of_unemployment_back_key")
    # if st.session_state.proof_of_unemployment_back is not None:
    #     if st.session_state.proof_of_unemployment_back not in st.session_state.files:
    #         st.session_state.files.append(st.session_state.proof_of_unemployment_back)


    # # st.session_state.files["educational_certificates"] = st.file_uploader("Educational Certificates (if applicable)", type=["jpg", "png", "pdf"], accept_multiple_files=True, key="educational_certificates")
    st.session_state.educational_certificates = st.file_uploader("Educational Certificates (IF APPLICABLE)", type=["jpg", "png", "pdf", "docx"], key="educational_certificates_key")
    if st.session_state.educational_certificates is not None:
        if st.session_state.educational_certificates not in st.session_state.files:
            st.session_state.files.append(st.session_state.educational_certificates)

    # Allow users to proceed regardless of whether they uploaded documents
    if st.button("Next"):
        st.session_state.step = 10
        st.experimental_rerun()


elif st.session_state.step == 10:
    st.title("> 9: Enrolment Agreement")
    st.header("Course Details")
    st.write("COURSE NAME: Skills Bootcamp in Health and Social Care")
    st.write("DURATION: 12 weeks (3 days per week)")
    st.write("START DATE: [To be filled]")
    st.write("END DATE: [To be filled]")
    st.write("LOCATION: [To be filled]")
    
    st.title("Terms and Conditions")
    st.header("Commitment:")
    st.write("o The participant agrees to attend all scheduled sessions and actively participate in all course activities.")
    st.write("o The participant agrees to complete all assignments and assessments as required.")
    st.header("Code of Conduct:")
    st.write("o The participant agrees to adhere to the Skills Bootcamp code of conduct, which includes respectful behavior towards peers and instructors, punctuality, and adherence to health and safety regulations.")
    st.header("Support Services:")
    st.write("o The participant is entitled to access all support services provided by the Skills Bootcamp, including career advice, counseling, and additional learning support.")
    st.header("Privacy and Data Protection:")
    st.write("o The participant agrees to the collection and use of their personal data in accordance with Prevista Ltd's Privacy Notice and the data sharing policies of Surrey County Council.")
    st.header("Consent:")
    st.write("o The participant consents to being filmed for course development, evaluation, and marketing purposes.")
    st.write("o The participant consents to data sharing with Surrey County Council.")
    st.header("Evaluation and Feedback:")
    st.write("o The participant may be contacted for surveys or interviews as part of the program's evaluation.")
    st.header("Termination:")
    st.write("o Prevista Ltd. reserves the right to terminate this agreement if the participant fails to comply with the terms and conditions outlined in this agreement.")
    st.header("Declaration and Consent")
    st.text('''I confirm that the information provided in this form is accurate 
and complete to the best of my knowledge. I consent to the collection 
and use of my personal data in accordance with Prevista Ltd's Privacy 
Notice and the data sharing policies of Surrey County Council. I also 
consent to being filmed for course development, evaluation, and 
marketing purposes.''')
    
    
    st.write("Participant Signature")
    # st.session_state.signature = st.text_input("Participant Signature")
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )
    st.session_state.signature = canvas_result.image_data
    st.session_state.date = st.date_input("Date", help="Choose a date", format='DD/MM/YYYY')

    if st.button("Submit"):
        if is_signature_drawn(st.session_state.signature) and st.session_state.date:
            st.session_state.submission_done = True
            st.session_state.step = 11
            st.experimental_rerun()
        else:
            st.warning("Please provide your signature before submitting.")

elif st.session_state.step == 11:
    st.write("**Thank you for completing the enrollment form!**")
    st.write("We will process your application and get in touch with you soon.")

# ####################################################################################################################################

# Generate and save the document if form is submitted
if st.session_state.submission_done:
    # doc = Document()
    # doc.add_heading('Skills Bootcamp Enrollment and Registration Document', 0)
    # doc.add_paragraph('Provider: Prevista Ltd. | Sponsor: Surrey County Council | Website: www.prevista.co.uk')
    # doc.add_paragraph('________________________________________')
    # doc.add_heading('Application and Enrollment Form', level=1)
    
    # # Adding content based on user input
    # doc.add_heading('1. Personal Information', level=2)
    # doc.add_paragraph(f"First Name: {st.session_state.first_name}")
    # doc.add_paragraph(f"Sir Name: {st.session_state.sir_name}")

    # DoB = st.session_state.dob.strftime('%d-%m-%Y')
    # doc.add_paragraph(f"Date of Birth: {DoB}")

    # doc.add_paragraph(f"Address: {st.session_state.address}")
    # doc.add_paragraph(f"City: {st.session_state.city}")
    # doc.add_paragraph(f"Postcode: {st.session_state.postcode}")
    # doc.add_paragraph(f"Phone Number: {st.session_state.phone_number}")
    # doc.add_paragraph(f"Email Address: {st.session_state.email}")
    # doc.add_paragraph(f"National Insurance Number: {st.session_state.ni_number}")
    # doc.add_paragraph(f"Emergency Contact Name: {st.session_state.emergency_contact_name}")
    # doc.add_paragraph(f"Emergency Contact Relationship: {st.session_state.emergency_contact_relationship}")
    # doc.add_paragraph(f"Emergency Contact Phone Number: {st.session_state.emergency_contact_phone}")
    # doc.add_paragraph(f"Emergency Contact Email Address: {st.session_state.emergency_contact_email}")
    
    # doc.add_heading('2. Educational Background', level=2)
    # doc.add_paragraph(f"Highest Level of Education: {st.session_state.highest_education}")
    # doc.add_paragraph(f"Institution Name: {st.session_state.institution_name}")
    # doc.add_paragraph(f"Year of Completion: {st.session_state.year_of_completion}")
    # doc.add_paragraph(f"COMPLETED LEVEL 6 OR ABOVE?: {st.session_state.completed_level_6_or_above}")
    # doc.add_paragraph(f"IF YES, WHAT SUBJECT: {st.session_state.subject}")
    # doc.add_paragraph(f"DATE APPLIED TO JOIN SKILLS BOOTCAMP? (MM/YYYY): {st.session_state.date_applied}")
    
    # doc.add_heading('3. Employment Status', level=2)
    # doc.add_paragraph(f"Current Employment Status: {st.session_state.current_employment_status}")
    # doc.add_paragraph(f"Previous Employer Name: {st.session_state.previous_employment_employer_name}")
    # doc.add_paragraph(f"Previous Job Title: {st.session_state.previous_employment_job_title}")
    # doc.add_paragraph(f"Duration of Employment: {st.session_state.previous_employment_duration}")

    
    # doc.add_heading('4. Additional Information', level=2)
    # doc.add_paragraph(f"Disabilities or Learning Difficulties: {st.session_state.disabilities}")
    # if st.session_state.disabilities == "Yes":
    #     doc.add_paragraph(f"Details: {st.session_state.disabilities_details}")
    # doc.add_paragraph(f"Receiving Benefits: {st.session_state.benefits}")
    # if st.session_state.benefits == "Yes":
    #     doc.add_paragraph(f"Details: {st.session_state.benefits_details}")
    
    # doc.add_heading('5. Statement of Interest', level=2)
    # doc.add_paragraph(f"Statement of Interest: {st.session_state.statement_of_interest}")
    
    # doc.add_heading('6. CEIAG (Careers Education, Information, Advice, and Guidance)', level=2)
    # doc.add_paragraph(f"Career Goals: {st.session_state.career_goals}")
    # doc.add_paragraph(f"Preferred Industry: {st.session_state.preferred_industry}")
    # doc.add_paragraph(f"Support Needed: {st.session_state.support_needed}")
    
    # doc.add_heading('7. Eligibility Criteria', level=2)
    # doc.add_paragraph(f"Age Verified: {'Yes' if st.session_state.age_verified else 'No'}")
    # doc.add_paragraph(f"Residency Verified: {'Yes' if st.session_state.residency_verified else 'No'}")
    # # doc.add_paragraph(f"Unemployment Verified: {'Yes' if st.session_state.unemployment_verified else 'No'}")
    # doc.add_paragraph(f"Education Verified: {'Yes' if st.session_state.education_verified else 'No'}")
    # doc.add_paragraph(f"Commitment Verified: {'Yes' if st.session_state.commitment_verified else 'No'}")
    
    # # doc.add_heading('8. Supporting Documents', level=2)
    # # doc.add_paragraph(f"Documents Verified: {'Yes' if st.session_state.documents_verified else 'No'}")
    
    # doc.add_heading('8. Enrolment Agreement', level=2)
    # # Save the signature image if available
    # if st.session_state.signature is not None:
    #     # Convert numpy array to PIL image
    #     image_data = st.session_state.signature
    #     image = Image.fromarray(image_data.astype(np.uint8))  # Ensure correct data type
        
    #     # Save the image to an in-memory file
    #     image_stream = io.BytesIO()
    #     image.save(image_stream, format='PNG')
    #     image_stream.seek(0)
        
    #     # Add image to docx
    #     doc.add_picture(image_stream, width=Inches(2))

    # submission_date = st.session_state.date.strftime('%d-%m-%Y')
    # doc.add_paragraph(f"Date: {submission_date}")

    # # Save the document
    # doc_path = f"Skills_Bootcamp_Submission_{st.session_state.first_name}_{st.session_state.sir_name}.docx"
    # doc.save(doc_path)

    # FILL TEMPLATE:
    placeholder_values = {
        'ph1': st.session_state.first_name,
        'ph2': st.session_state.sir_name,
        'ph3': st.session_state.ni_number,
        'ph4': st.session_state.current_employer_postcode,
        'ph5': st.session_state.email,
        'ph6': st.session_state.phone_number,
        # 'ph7': st.session_state.placeholder_7,
        # 'ph8': st.session_state.placeholder_8,
        # 'ph9': st.session_state.placeholder_9,
        'ph10': st.session_state.highest_education,
        'ph11': st.session_state.institution_name,
        'ph12': st.session_state.year_of_completion,
        'ph13': st.session_state.completed_level_6_or_above,
        'ph14': st.session_state.subject,
        'ph15': st.session_state.date_applied,
        # 'ph16': st.session_state.placeholder_16,
        # 'ph17': st.session_state.placeholder_17,
        # 'ph18': st.session_state.placeholder_18,
        # 'ph19': st.session_state.placeholder_19,
        'ph20': st.session_state.current_employment_status,
        'ph21': st.session_state.current_employer_name,
        'ph22': st.session_state.current_employer_postcode,
        'ph23': st.session_state.applied_through_employer,
        'ph24': st.session_state.hours_worked_per_week,
        'ph25': st.session_state.estimated_salary,
        'ph26': st.session_state.salary_frequency,
        'ph27': st.session_state.contract_type,
        'ph28': st.session_state.continue_working,
        'ph29': st.session_state.main_work,
        'ph30': st.session_state.previous_employment_job_title,
        'ph31': st.session_state.previous_employment_industry,
        
        'ph32': st.session_state.claiming_universal_credit,
        'ph33': st.session_state.caring_responsibilities,
        'ph34': st.session_state.dob,
        'ph35': st.session_state.gender,
        'ph36': st.session_state.disabilities_health_condition,
        'ph37': st.session_state.ethnicity,

        'ph38': st.session_state.emergency_contact_name,
        'ph39': st.session_state.emergency_contact_relationship,
        'ph40': st.session_state.emergency_contact_phone,
        'ph41': st.session_state.emergency_contact_email,

        # 'ph42': st.session_state.placeholder_42,
        # 'ph43': st.session_state.placeholder_43,
        # 'ph44': st.session_state.placeholder_44,
        'ph45': st.session_state.statement_of_interest,
        'ph46': st.session_state.career_goals,
        'ph47': st.session_state.preferred_industry,
        'ph48': st.session_state.support_needed,

        'ph50': date.today(),
    }

    template_file = "ph_skills_bootcamp.docx"
    modified_file = f"SkillsBootcamp_Form_Submission_{st.session_state.first_name}_{st.session_state.sir_name}.docx"

    signature_path = 'signature_image.png'
    signature_image = PILImage.fromarray(
        st.session_state.signature.astype('uint8'), 'RGBA')
    signature_image.save(signature_path)

    replace_placeholders(template_file, modified_file, placeholder_values, signature_path)



# Email
    # Sender email credentials
    # Credentials: Streamlit host st.secrets
    sender_email = st.secrets["sender_email"]
    sender_password = st.secrets["sender_password"]

    # Credentials: Local env
    # load_dotenv()                                     # uncomment import of this library!
    # sender_email = os.getenv('EMAIL')
    # sender_password = os.getenv('PASSWORD')
    # team_email = [sender_email]
    team_email = ['muhammadoa@prevista.co.uk']
    # receiver_email = sender_email
    # receiver_email = 'mohamedr@prevista.co.uk'

    learner_email = [st.session_state.email]
    
    subject_team = f"Skills_Bootcamp Name: {st.session_state.first_name}_{st.session_state.sir_name} Submission Date: {date.today()}"
    body_team = "Prevista Skills Bootcamp Form submitted. Please find attached files."

    subject_learner = "Thank You for Your Interest in The Skills Bootcamp!"
    body_learner = f"""
    <html>
    <body>
        <p>Dear {st.session_state.first_name} {st.session_state.sir_name},</p>

        <p>Thank you for expressing your interest in Bootcamp at PREVISTA. We are excited to guide you through the next steps of the enrollment process.</p>

        <p><strong>What’s Next?</strong></p>
        <ol>
            <li><strong>Enrollment Communication:</strong> One of our representatives will be contacting you within the next few days to complete your enrollment. Please keep an eye out for our message to finalize your registration details.</li>
            <li><strong>Course Start Date:</strong> Once your enrollment is confirmed, we will send you the schedule for the course start date.</li>
            <li><strong>Orientation Session:</strong> You will be invited to an orientation session where you can learn more about the platform, meet your instructors, and connect with other learners.</li>
        </ol>

        <p>If you have any immediate questions, please feel free to reach out to us at [support email] or [support phone number].</p>

        <p>We look forward to speaking with you soon and welcoming you to our learning community!</p>

        <p>Best regards,</p>
        <p>Student Admissions Team<br>
        PREVISTA<br>
        PREPARING YOU TODAY FOR OPPORTUNITIES OF TOMORROW</p>
    </body>
    </html>
    """

    # Local file path
    local_file_path = modified_file

    # Send email to team with attachments
    if st.session_state.files or local_file_path:
        send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, st.session_state.files, local_file_path)
    
    # Send thank you email to learner
    send_email_with_attachments(sender_email, sender_password, learner_email, subject_learner, body_learner)

    update_doc()


# streamlit run app.py --server.port 8504
# Dev : https://linkedin.com/in/osamatech786
